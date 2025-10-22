import nbformat
from docx import Document

# Abre el archivo Jupyter Notebook
with open('C:/Users/pgimenezbarrera/Desktop/Patrícia GB/Codes/Application_Test_All_finale.ipynb', 'r', encoding='utf-8') as f:
    notebook_content = nbformat.read(f, as_version=4)  # Lee el archivo como un objeto NotebookNode

# Crear el documento Word
doc = Document()

# Iterar sobre las celdas del cuaderno Jupyter
for cell in notebook_content['cells']:
    if cell['cell_type'] == 'markdown':
        # Si es una celda de markdown (texto), agregarla al documento
        doc.add_paragraph(cell['source'])
    elif cell['cell_type'] == 'code':
        # Si es una celda de código, agregarla también
        doc.add_paragraph("Código: \n")
        doc.add_paragraph(cell['source'])

# Guardar el documento Word
doc.save('Application_Test_All_finale.docx')

print("El archivo Word ha sido creado correctamente.")





